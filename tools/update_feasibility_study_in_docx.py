# pyright: reportMissingImports=false, reportMissingTypeStubs=false, reportUnknownMemberType=false, reportPrivateUsage=false
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DOC_PATH = ROOT / "docs" / "Smart_Grocery_Project_Report.docx"
FALLBACK_PATH = ROOT / "docs" / "Smart_Grocery_Project_Report_Feasibility_Updated.docx"


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is None:
        return
    parent.remove(element)


def insert_paragraph_after(paragraph, text: str = "", style: str | None = None):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def format_body(paragraph, text: str):
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.first_line_indent = Inches(0.4)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def format_subheading(paragraph, text: str):
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_before = Pt(6)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)
    run.font.italic = True
    run.font.color.rgb = RGBColor(79, 129, 189)


def format_bullet(paragraph, text: str):
    paragraph.text = ""
    paragraph.style = "List Bullet"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    run.font.size = Pt(12)


def main():
    doc = Document(DOC_PATH)
    paragraphs = doc.paragraphs

    heading_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "1.4 Feasibility Study")
    next_heading_idx = next(
        i for i, p in enumerate(paragraphs[heading_idx + 1 :], start=heading_idx + 1)
        if p.text.strip().startswith("2.")
    )

    for idx in range(next_heading_idx - 1, heading_idx, -1):
        remove_paragraph(doc.paragraphs[idx])

    anchor = doc.paragraphs[heading_idx]

    content = [
        ("subheading", "Technical Feasibility:"),
        (
            "body",
            "The Smart Grocery system is technically feasible using modern full-stack web development technologies. "
            "The project is built using React and Vite for the frontend, Spring Boot for the backend, and MySQL for "
            "database management. These technologies are widely adopted, well-supported, and suitable for building a "
            "responsive and secure grocery management platform.",
        ),
        (
            "bullet",
            "Frontend: The frontend is built using React, Vite, Tailwind CSS, Axios, and React Router. These tools "
            "help create a responsive and interactive user interface for login, dashboard insights, inventory control, "
            "shopping suggestions, and admin management screens.",
        ),
        (
            "bullet",
            "Backend: The backend is powered by Spring Boot with Spring Security, validation, and JPA. It handles "
            "authentication, grocery CRUD operations, recommendation logic, expiry reminders, admin controls, and API routing "
            "in a structured and scalable manner.",
        ),
        (
            "bullet",
            "Database: MySQL is used for storing user accounts and grocery item records. Since it is a relational database "
            "management system, it can efficiently manage structured data and support future application growth.",
        ),
        (
            "bullet",
            "Security: Security is implemented using JWT-based authentication, password hashing, protected frontend routes, "
            "and role-based access control for normal users and administrators. These measures help protect user data and restrict "
            "sensitive operations to authorized roles.",
        ),
        ("subheading", "Operational Feasibility:"),
        (
            "body",
            "The Smart Grocery system is operationally feasible because it is easy to use and accessible through a web browser. "
            "Users can manage groceries, track stock levels, and view expiry reminders without requiring special training. The "
            "admin module also centralizes monitoring tasks such as user management, product review, stock control, and reports.",
        ),
        (
            "body",
            "Because the application is web-based, it can be accessed from standard computers and laptops in a simple local setup. "
            "This makes it practical for academic demonstration and convenient for day-to-day use.",
        ),
        ("subheading", "Economic Feasibility:"),
        (
            "body",
            "From an economic point of view, Smart Grocery is cost-effective because it uses open-source technologies such as React, "
            "Spring Boot, MySQL, and related JavaScript and Java libraries. This reduces licensing costs while still providing a strong "
            "technical foundation for development and deployment.",
        ),
        (
            "bullet",
            "Development Costs: The primary cost of development is the time and effort required for design, coding, testing, and documentation. "
            "Since the tools used in the project are open source, no expensive software licenses are required.",
        ),
        (
            "bullet",
            "Operational Costs: The operational cost is comparatively low because the system can run on a standard local machine for development "
            "and testing. If deployed later, affordable cloud or shared-hosting infrastructure can support the backend and database services.",
        ),
    ]

    current = anchor
    for kind, text in content:
        current = insert_paragraph_after(current)
        if kind == "subheading":
            format_subheading(current, text)
        elif kind == "bullet":
            format_bullet(current, text)
        else:
            format_body(current, text)

    try:
        doc.save(DOC_PATH)
        print(DOC_PATH)
    except PermissionError:
        doc.save(FALLBACK_PATH)
        print(FALLBACK_PATH)


if __name__ == "__main__":
    main()
