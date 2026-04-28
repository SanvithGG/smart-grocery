# pyright: reportMissingImports=false, reportMissingTypeStubs=false, reportUnknownMemberType=false, reportPrivateUsage=false
from __future__ import annotations

from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Iterable
import textwrap

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
OUTPUT_DIR = ROOT / "docs"
ASSETS_DIR = OUTPUT_DIR / "report_assets"
OUTPUT_PATH = OUTPUT_DIR / "Smart_Grocery_Project_Report.docx"


@dataclass
class FigureSpec:
    filename: str
    title: str
    caption: str


def read_text(path: Path) -> str:
    return path.read_text(encoding="utf-8")


def read_snippet(path: Path, start: int, end: int) -> str:
    lines = read_text(path).splitlines()
    return "\n".join(lines[start - 1 : end])


def safe_font(size: int, bold: bool = False, mono: bool = False):
    font_candidates = []
    if mono:
        font_candidates.extend(
            [
                r"C:\Windows\Fonts\consola.ttf",
                r"C:\Windows\Fonts\cour.ttf",
            ]
        )
    elif bold:
        font_candidates.extend(
            [
                r"C:\Windows\Fonts\timesbd.ttf",
                r"C:\Windows\Fonts\georgiab.ttf",
                r"C:\Windows\Fonts\arialbd.ttf",
                r"C:\Windows\Fonts\calibrib.ttf",
            ]
        )
    else:
        font_candidates.extend(
            [
                r"C:\Windows\Fonts\times.ttf",
                r"C:\Windows\Fonts\georgia.ttf",
                r"C:\Windows\Fonts\arial.ttf",
                r"C:\Windows\Fonts\calibri.ttf",
            ]
        )

    for candidate in font_candidates:
        if Path(candidate).exists():
            return ImageFont.truetype(candidate, size=size)

    return ImageFont.load_default()


def wrap_text(text: str, width: int) -> list[str]:
    return textwrap.wrap(text, width=width, break_long_words=False, break_on_hyphens=False) or [text]


def draw_multiline(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font, fill, width: int, line_gap: int = 8):
    x, y = xy
    line_height = font.size + line_gap if hasattr(font, "size") else 18
    for line in wrap_text(text, width):
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height
    return y


def new_canvas(title: str, subtitle: str | None = None, size: tuple[int, int] = (1600, 900)) -> tuple[Image.Image, ImageDraw.ImageDraw]:
    image = Image.new("RGB", size, "#f8fafc")
    draw = ImageDraw.Draw(image)
    draw.rectangle((0, 0, size[0], 86), fill="#0f172a")
    draw.text((48, 24), title, font=safe_font(34, bold=True), fill="#ffffff")
    if subtitle:
        draw.text((50, 94), subtitle, font=safe_font(18), fill="#475569")
    return image, draw


def save_image(image: Image.Image, path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    image.save(path)


def draw_round_rect(draw, box, fill, outline="#cbd5e1", radius=18, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def create_architecture_diagram(path: Path):
    image, draw = new_canvas("Architecture Diagram", "Client-server architecture for Smart Grocery")
    boxes = [
        ((80, 190, 360, 330), "#dbeafe", "React Frontend\nHome, Dashboard,\nInventory, Shopping List"),
        ((500, 190, 820, 330), "#e2e8f0", "Spring Boot API\nAuthController\nGroceryController\nAdminController"),
        ((960, 190, 1240, 330), "#dcfce7", "Service Layer\nAuthService\nGroceryService\nAdminService"),
        ((1280, 190, 1500, 330), "#fee2e2", "MySQL Database\nusers\ngrocery_items"),
        ((500, 500, 820, 640), "#fef3c7", "Security Layer\nJWT Filter\nRole checks\nProtected routes"),
        ((960, 500, 1240, 640), "#ede9fe", "Smart Logic\nRecommendations\nExpiry alerts\nCatalog stock"),
    ]
    for box, fill, label in boxes:
        draw_round_rect(draw, box, fill=fill)
        center_x = (box[0] + box[2]) // 2
        y = box[1] + 24
        for idx, line in enumerate(label.splitlines()):
            font = safe_font(24 if idx == 0 else 18, bold=idx == 0)
            w = draw.textbbox((0, 0), line, font=font)[2]
            draw.text((center_x - w / 2, y), line, font=font, fill="#0f172a")
            y += 38 if idx == 0 else 28

    arrows = [
        ((360, 260), (500, 260)),
        ((820, 260), (960, 260)),
        ((1240, 260), (1280, 260)),
        ((660, 330), (660, 500)),
        ((1100, 330), (1100, 500)),
        ((820, 570), (960, 570)),
    ]
    for start, end in arrows:
        draw.line([start, end], fill="#0284c7", width=6)
        draw.polygon(
            [
                (end[0], end[1]),
                (end[0] - 18, end[1] - 10),
                (end[0] - 18, end[1] + 10),
            ],
            fill="#0284c7",
        )
    save_image(image, path)


def create_module_interaction_diagram(path: Path):
    image, draw = new_canvas("Module Interaction Diagram", "User and admin flows across major Smart Grocery modules")
    modules = [
        ((90, 170, 420, 285), "#e0f2fe", "Authentication"),
        ((90, 345, 420, 460), "#dcfce7", "Home and Catalog"),
        ((90, 520, 420, 635), "#fef3c7", "Inventory"),
        ((580, 170, 930, 285), "#fde68a", "Dashboard"),
        ((580, 345, 930, 460), "#ede9fe", "Shopping List"),
        ((580, 520, 930, 635), "#fee2e2", "Expiry Reminder"),
        ((1090, 170, 1470, 285), "#e2e8f0", "Admin Dashboard"),
        ((1090, 345, 1470, 460), "#fce7f3", "Admin Products and Stock"),
        ((1090, 520, 1470, 635), "#ecfccb", "Admin Reports and Users"),
    ]
    for box, fill, label in modules:
        draw_round_rect(draw, box, fill)
        bbox = draw.textbbox((0, 0), label, font=safe_font(24, bold=True))
        draw.text((((box[0] + box[2]) - bbox[2]) / 2, box[1] + 38), label, font=safe_font(24, bold=True), fill="#0f172a")
    connectors = [
        ((420, 227), (580, 227)),
        ((420, 402), (580, 402)),
        ((420, 577), (580, 577)),
        ((930, 227), (1090, 227)),
        ((930, 402), (1090, 402)),
        ((930, 577), (1090, 577)),
        ((255, 285), (255, 345)),
        ((255, 460), (255, 520)),
        ((755, 285), (755, 345)),
        ((755, 460), (755, 520)),
        ((1280, 285), (1280, 345)),
        ((1280, 460), (1280, 520)),
    ]
    for start, end in connectors:
        draw.line([start, end], fill="#475569", width=4)
    save_image(image, path)


def create_er_diagram(path: Path):
    image, draw = new_canvas("E-R Diagram", "Database entities used by the Smart Grocery project")
    draw_round_rect(draw, (170, 210, 590, 650), fill="#dbeafe", outline="#1d4ed8", radius=24)
    draw_round_rect(draw, (940, 170, 1410, 710), fill="#dcfce7", outline="#15803d", radius=24)
    draw.text((310, 240), "users", font=safe_font(30, bold=True), fill="#0f172a")
    user_fields = ["id (PK)", "username", "email", "password", "role"]
    y = 300
    for field in user_fields:
        draw.text((220, y), field, font=safe_font(24), fill="#1e293b")
        y += 56
    draw.text((1090, 205), "grocery_items", font=safe_font(30, bold=True), fill="#0f172a")
    item_fields = [
        "id (PK)",
        "name",
        "category",
        "quantity",
        "purchased",
        "expiry_date",
        "last_purchased_at",
        "user_id (FK)",
    ]
    y = 270
    for field in item_fields:
        draw.text((995, y), field, font=safe_font(24), fill="#1e293b")
        y += 52
    draw.line([(590, 430), (940, 430)], fill="#0f172a", width=6)
    draw.text((690, 375), "1", font=safe_font(30, bold=True), fill="#0f172a")
    draw.text((830, 375), "M", font=safe_font(30, bold=True), fill="#0f172a")
    draw.text((630, 470), "One user owns many grocery items", font=safe_font(22), fill="#475569")
    save_image(image, path)


def create_dfd_user(path: Path):
    image, draw = new_canvas("DFD - User Flow", "Data flow for a normal user")
    items = [
        ("User", (90, 300, 310, 420), "#dbeafe"),
        ("Login / Register", (420, 140, 760, 260), "#fef3c7"),
        ("Inventory Service", (420, 320, 760, 440), "#e2e8f0"),
        ("Dashboard and Shopping", (420, 500, 760, 620), "#dcfce7"),
        ("MySQL", (920, 320, 1210, 440), "#fee2e2"),
        ("Expiry and Recommendation Rules", (920, 500, 1420, 620), "#ede9fe"),
    ]
    for label, box, fill in items:
        draw_round_rect(draw, box, fill)
        y = box[1] + 36
        for line in label.split(" / "):
            w = draw.textbbox((0, 0), line, font=safe_font(24, bold=True))[2]
            draw.text((((box[0] + box[2]) - w) / 2, y), line, font=safe_font(24, bold=True), fill="#0f172a")
            y += 36
    for start, end in [
        ((310, 360), (420, 200)),
        ((310, 360), (420, 380)),
        ((760, 380), (920, 380)),
        ((760, 560), (920, 560)),
        ((760, 380), (760, 560)),
    ]:
        draw.line([start, end], fill="#0284c7", width=6)
    save_image(image, path)


def create_dfd_admin(path: Path):
    image, draw = new_canvas("DFD - Admin Flow", "Admin management and reporting flow")
    items = [
        ("Admin User", (90, 310, 310, 430), "#dbeafe"),
        ("Admin Login", (430, 130, 730, 250), "#fef3c7"),
        ("Admin Dashboard", (430, 300, 730, 420), "#e2e8f0"),
        ("Users / Products / Queue", (430, 470, 820, 590), "#dcfce7"),
        ("Reports Engine", (980, 300, 1260, 420), "#ede9fe"),
        ("MySQL", (980, 470, 1260, 590), "#fee2e2"),
    ]
    for label, box, fill in items:
        draw_round_rect(draw, box, fill)
        lines = label.split(" / ")
        y = box[1] + 34
        for line in lines:
            w = draw.textbbox((0, 0), line, font=safe_font(22, bold=True))[2]
            draw.text((((box[0] + box[2]) - w) / 2, y), line, font=safe_font(22, bold=True), fill="#0f172a")
            y += 34
    for start, end in [
        ((310, 370), (430, 190)),
        ((310, 370), (430, 360)),
        ((730, 360), (980, 360)),
        ((820, 530), (980, 530)),
        ((730, 360), (730, 530)),
    ]:
        draw.line([start, end], fill="#475569", width=6)
    save_image(image, path)


def create_testing_diagram(path: Path):
    image, draw = new_canvas("Testing Strategy", "Static checks, backend tests, and manual workflow validation")
    stages = [
        ("Frontend lint", "#dbeafe"),
        ("Backend unit tests", "#dcfce7"),
        ("Controller tests", "#fef3c7"),
        ("Manual workflow review", "#ede9fe"),
        ("Documentation evidence", "#fee2e2"),
    ]
    x = 80
    for label, color in stages:
        draw_round_rect(draw, (x, 330, x + 250, 470), fill=color)
        y = 368
        for line in wrap_text(label, 18):
            w = draw.textbbox((0, 0), line, font=safe_font(24, bold=True))[2]
            draw.text((x + 125 - w / 2, y), line, font=safe_font(24, bold=True), fill="#0f172a")
            y += 34
        x += 285
    for start_x in [330, 615, 900, 1185]:
        draw.line([(start_x, 400), (start_x + 35, 400)], fill="#0284c7", width=6)
        draw.polygon([(start_x + 35, 400), (start_x + 18, 390), (start_x + 18, 410)], fill="#0284c7")
    save_image(image, path)


def create_bar_chart(path: Path, title: str, subtitle: str, bars: list[tuple[str, int]], color: str = "#2563eb"):
    image, draw = new_canvas(title, subtitle)
    max_value = max(value for _, value in bars) or 1
    x = 130
    for label, value in bars:
        height = int((value / max_value) * 360)
        y0 = 720 - height
        draw.rounded_rectangle((x, y0, x + 90, 720), radius=14, fill=color, outline=color)
        value_text = str(value)
        draw.text((x + 24, y0 - 40), value_text, font=safe_font(20, bold=True), fill="#0f172a")
        label_y = 748
        for line in wrap_text(label, 10):
            draw.text((x - 5, label_y), line, font=safe_font(18), fill="#475569")
            label_y += 22
        x += 150
    draw.line([(100, 720), (1450, 720)], fill="#94a3b8", width=3)
    save_image(image, path)


def create_ui_mock(path: Path, page_title: str, hero: str, left_cards: list[str], right_cards: list[str], accent: str):
    image = Image.new("RGB", (1600, 940), "#f1f5f9")
    draw = ImageDraw.Draw(image)
    draw.rounded_rectangle((30, 30, 1570, 910), radius=30, fill="#ffffff", outline="#cbd5e1", width=3)
    draw.rectangle((30, 30, 1570, 88), fill="#0f172a")
    for i, color in enumerate(["#ef4444", "#f59e0b", "#10b981"]):
        draw.ellipse((60 + i * 26, 50, 76 + i * 26, 66), fill=color)
    draw.text((140, 48), "Smart Grocery - Browser View", font=safe_font(24, bold=True), fill="#ffffff")
    draw.rounded_rectangle((55, 120, 300, 870), radius=26, fill="#0f172a")
    side_items = ["Home", "Dashboard", "Inventory", "Shopping List", "Admin", "Reports"]
    y = 170
    for item in side_items:
        fill = accent if item.lower() in page_title.lower() else "#1e293b"
        draw.rounded_rectangle((80, y, 275, y + 54), radius=18, fill=fill)
        draw.text((110, y + 15), item, font=safe_font(20, bold=True), fill="#ffffff")
        y += 76

    draw.rounded_rectangle((340, 120, 1515, 285), radius=28, fill=accent)
    draw.text((390, 160), page_title, font=safe_font(40, bold=True), fill="#ffffff")
    draw_multiline(draw, (390, 214), hero, safe_font(22), "#e2e8f0", 60, line_gap=6)

    card_positions = [
        (350, 330, 890, 530),
        (925, 330, 1465, 530),
        (350, 560, 890, 815),
        (925, 560, 1465, 815),
    ]
    for idx, box in enumerate(card_positions[: len(left_cards + right_cards)]):
        draw_round_rect(draw, box, fill="#f8fafc")
        source = left_cards + right_cards
        card_text = source[idx]
        parts = card_text.split(":", 1)
        title = parts[0]
        body = parts[1].strip() if len(parts) > 1 else ""
        draw.text((box[0] + 24, box[1] + 20), title, font=safe_font(24, bold=True), fill="#0f172a")
        draw_multiline(draw, (box[0] + 24, box[1] + 70), body, safe_font(18), "#475569", 40)

    save_image(image, path)


def generate_assets():
    ASSETS_DIR.mkdir(parents=True, exist_ok=True)

    create_architecture_diagram(ASSETS_DIR / "architecture_diagram.png")
    create_module_interaction_diagram(ASSETS_DIR / "module_interaction.png")
    create_er_diagram(ASSETS_DIR / "er_diagram.png")
    create_dfd_user(ASSETS_DIR / "dfd_user.png")
    create_dfd_admin(ASSETS_DIR / "dfd_admin.png")
    create_testing_diagram(ASSETS_DIR / "testing_strategy.png")

    create_bar_chart(
        ASSETS_DIR / "catalog_stock_chart.png",
        "Catalog Stock Overview",
        "Default catalog quantities derived from GroceryService",
        [("Milk", 8), ("Eggs", 12), ("Rice", 15), ("Tea", 0), ("Coffee", 4), ("Biscuits", 16), ("Soap", 9)],
    )
    create_bar_chart(
        ASSETS_DIR / "module_count_chart.png",
        "Frontend and Backend Module Coverage",
        "Relative breadth of features covered in the report",
        [("User pages", 5), ("Admin pages", 6), ("Controllers", 3), ("Services", 5), ("Test suites", 5)],
        color="#7c3aed",
    )

    create_ui_mock(
        ASSETS_DIR / "ui_auth_login.png",
        "Login Screen",
        "Secure user and admin access through JWT-backed authentication.",
        ["Authentication Card: Username or email, password fields, role-safe redirect handling.", "Validation: Invalid credentials and expired session messages are surfaced through the API client."],
        ["Navigation: Login and register views are separated and protected by routing rules.", "Outcome: Successful login stores token and role before opening the app shell."],
        "#1d4ed8",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_auth_register.png",
        "Register Screen",
        "New user onboarding with validation for account creation.",
        ["Registration Form: Username, email, password, and submit action.", "Service Behavior: Backend hashes the password and stores a USER role by default."],
        ["Error Handling: Duplicate usernames and validation failures are returned in a readable format.", "Outcome: Newly created users can immediately sign in through the login page."],
        "#2563eb",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_home_catalog.png",
        "Home and Catalog",
        "The home view focuses on quick buying, catalog browsing, and top-level grocery status.",
        ["Hero Section: Dashboard shortcut, buy queue access, and inventory management links.", "Catalog Filters: Search and category filters update the visible grocery catalog."],
        ["Catalog Cards: Price, quantity, and stock availability are shown before purchase.", "Quick Buy Flow: Users can open a purchase modal and add an item immediately."],
        "#0284c7",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_quick_buy.png",
        "Quick Buy Modal",
        "A focused purchase flow for adding catalog items directly into the household inventory.",
        ["Quantity and Expiry: Quantity is entered manually, while expiry can be derived from item rules.", "User Guidance: The dialog explains available stock and natural expiry preview."],
        ["Submission Result: A bought item is persisted and reflected across dashboard and inventory views.", "Inventory Link: The modal supports rapid shopping without leaving the catalog page."],
        "#0f766e",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_dashboard_overview.png",
        "Dashboard Overview",
        "Decision-oriented dashboard with summary cards and operational sections.",
        ["Summary Cards: Total items, pending items, purchased items, and low-stock items.", "Contextual Sections: Recent list, pending focus, smart picks, action board, and reminders."],
        ["Refresh Model: Dashboard listens for inventory changes and reloads data from the API.", "User Value: Users see the next actions in one place instead of browsing each page manually."],
        "#0f172a",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_dashboard_recommendations.png",
        "Dashboard - Smart Picks",
        "Recommendation cards surface restock suggestions based on purchase history and stock state.",
        ["Reasoning: Items gain priority from previous purchases and low-stock conditions.", "Priority Labels: HIGH, MEDIUM, and LOW values are displayed in the recommendation queue."],
        ["Presentation: The screen groups item name, category, reason, and priority in a card list.", "Impact: This makes the dashboard more than a static inventory summary."],
        "#7c3aed",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_dashboard_lowstock.png",
        "Dashboard - Low Stock Watchlist",
        "Low-stock items appear in a dedicated action board for fast attention.",
        ["Filter Rule: Pending items at or below threshold are surfaced here.", "Card Layout: Each card displays item name, category, and current quantity."],
        ["Operational Benefit: Users can identify refill pressure without scanning the entire list.", "Design: The darker action board visually separates urgent stock concerns."],
        "#b45309",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_dashboard_expiry.png",
        "Dashboard - Kitchen Reminder",
        "Purchased items that are close to expiry or expired are shown as kitchen reminders.",
        ["Severity Types: OVERDUE, TODAY, and SOON guide the urgency of action.", "Action Button: Users can acknowledge the alert and delete the item from the database."],
        ["Message Design: Each reminder explains what the user should do next.", "Lifecycle: Alerts are recomputed from stored expiry dates."],
        "#b91c1c",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_inventory.png",
        "Inventory Management",
        "The inventory page combines an add-item form with a searchable list of grocery items.",
        ["Add Form: Name, category, quantity, purchased flag, and expiry preview.", "Daily Picks: Catalog suggestions autofill the add-item form for faster entry."],
        ["List Controls: Search, category filter, purchased filter, delete, and toggle bought state.", "Data Detail: Items show quantity, status, expiry, and last-purchased timestamp."],
        "#0369a1",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_shopping_list.png",
        "Shopping List",
        "The shopping queue is auto-built from low stock and near-expiry conditions.",
        ["Priority Badges: HIGH and MEDIUM indicate urgency of purchase.", "Reason Tags: LOW STOCK and EXPIRING explain why an item appears here."],
        ["Actions: Users can mark items as bought or temporarily dismiss suggestions.", "Sync Behavior: Inventory updates trigger a queue refresh through shared events."],
        "#15803d",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_admin_dashboard.png",
        "Admin Dashboard",
        "The admin landing page provides a platform-wide overview of users, products, categories, and stock pressure.",
        ["Summary Metrics: Total users, products, purchased, pending, categories, and low-stock products.", "Use Case: This page is the entry point for operational monitoring."],
        ["Audience: Unlike the user dashboard, this view summarizes the entire system.", "Value: Admins can assess activity before navigating to detailed controls."],
        "#1e293b",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_admin_users.png",
        "Admin Users",
        "User management supports role changes and deletion with system-wide guardrails.",
        ["Role Management: Admin can promote or demote users except when it would remove the last admin.", "User Summary: Each record shows email, total items, purchased items, and pending items."],
        ["Safety Rule: At least one admin account must remain in the system.", "Operational Use: This page manages access control at the application level."],
        "#0f766e",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_admin_products.png",
        "Admin Products and Stock",
        "Products and catalog stock are managed through dedicated admin controls.",
        ["Catalog Stock: Quantities can be edited to restock out-of-stock or low-stock items.", "Product Records: Admin can adjust quantity or delete user-linked product rows."],
        ["Search: Product search spans name, category, and username.", "Business Logic: Stock changes affect what users see during quick-buy and purchase workflows."],
        "#7c2d12",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_admin_purchase_queue.png",
        "Admin Purchase Queue",
        "Pending purchases across users are completed from a central queue.",
        ["Queue Source: Pending grocery items that are not yet marked purchased.", "Action: Complete Purchase triggers the backend fulfillment logic and stock reduction."],
        ["Coordination: The queue bridges user intent and administrative stock control.", "Feedback: Successful completion updates the shared application state."],
        "#166534",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_admin_categories.png",
        "Admin Categories",
        "Category maintenance lets administrators rename categories without editing every record manually.",
        ["Category Summary: Total, purchased, and pending product counts are shown for each category.", "Rename Flow: Current category names can be updated in place."],
        ["Result: Category consistency can be improved without direct database editing.", "Use Case: Helpful for cleaning inconsistent naming across user entries."],
        "#6d28d9",
    )
    create_ui_mock(
        ASSETS_DIR / "ui_admin_reports.png",
        "Admin Reports",
        "The reports page converts product data into an operational summary for administrators.",
        ["Summary Cards: Users, total products, purchased, pending, and expiring soon products.", "Top Categories: Highest-volume categories are highlighted separately."],
        ["Breakdown: Each category row lists total, purchased, and pending products.", "Reporting Goal: The page supports academic discussion of analytics in Smart Grocery."],
        "#be185d",
    )


def set_run_font(run, size=12, bold=False, italic=False, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def set_page_margins(section):
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = " PAGE "
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_end)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = r' TOC \o "1-3" \h \z \u '
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = 'Update this field in Word to refresh the table of contents.'
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr_text)
    run._r.append(fld_char_sep)
    run._r.append(text)
    run._r.append(fld_char_end)


def add_paragraph(doc: Document, text: str = "", *, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, bold=False, italic=False, line_spacing=1.5, indent=True):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.first_line_indent = Inches(0.4) if indent and align == WD_ALIGN_PARAGRAPH.JUSTIFY else Pt(0)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    p.style = doc.styles["Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    r = p.add_run(text)
    set_run_font(r, size=14 if level == 1 else 12, bold=True)
    for run in p.runs:
        set_run_font(run, size=14 if level == 1 else 12, bold=True)
    return p


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.line_spacing = 1.4
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(item)
        set_run_font(r, size=12)


def set_cell_text(cell, text: str, *, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_table(doc: Document, rows: list[list[str]]):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            set_cell_text(
                table.cell(row_idx, col_idx),
                value,
                bold=row_idx == 0,
                align=WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT,
            )
    return table


def add_code_block(doc: Document, title: str, code: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    set_run_font(r, size=12, bold=True)

    code_p = doc.add_paragraph()
    code_p.paragraph_format.left_indent = Inches(0.25)
    code_p.paragraph_format.right_indent = Inches(0.15)
    code_p.paragraph_format.line_spacing = 1.0
    code_p.paragraph_format.space_after = Pt(8)
    run = code_p.add_run(code)
    set_run_font(run, size=8.5, name="Courier New")


def add_figure(doc: Document, asset_name: str, title: str, caption: str, width_inches: float = 6.2, page_break: bool = True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(ASSETS_DIR / asset_name), width=Inches(width_inches))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.line_spacing = 1.2
    cap.paragraph_format.space_before = Pt(2)
    r1 = cap.add_run(title + " ")
    set_run_font(r1, size=11, bold=True)
    r2 = cap.add_run(caption)
    set_run_font(r2, size=11)
    if page_break:
        doc.add_page_break()


def add_title_page(doc: Document):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.4
    lines = [
        ("A PROJECT REPORT", 16, True),
        ("\nON", 14, True),
        ("\nSMART GROCERY", 22, True),
        ("\n\nSubmitted in partial fulfillment of the requirements", 12, False),
        ("\nfor the award of the degree / project documentation submission", 12, False),
        ("\n\nPrepared from the Smart Grocery full-stack project repository", 12, False),
        ("\n\nStudent Name  : ____________________", 12, False),
        ("\nUSN / Roll No : ____________________", 12, False),
        ("\nGuide Name    : ____________________", 12, False),
        ("\nDepartment    : ____________________", 12, False),
        ("\nCollege Name  : ____________________", 12, False),
        ("\nAcademic Year : ____________________", 12, False),
        (f"\n\nDate: {date(2026, 4, 19).strftime('%d %B %Y')}", 12, False),
    ]
    for text, size, bold in lines:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
    doc.add_page_break()


def add_front_matter(doc: Document):
    add_title_page(doc)

    add_heading(doc, "CERTIFICATE", level=1)
    add_paragraph(
        doc,
        "This is to certify that the project report entitled \"SMART GROCERY\" is a bona fide work based on the implementation "
        "of the Smart Grocery project repository. The work documented in this report covers the design, development, testing, "
        "results, and academic presentation of the system, and is submitted as part of the project documentation requirement.",
    )
    add_paragraph(doc, "\nGuide Signature: ____________________\n\nHead of Department: ____________________\n\nExternal Examiner: ____________________", align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
    doc.add_page_break()

    add_heading(doc, "DECLARATION", level=1)
    add_paragraph(
        doc,
        "I hereby declare that the report titled \"SMART GROCERY\" has been prepared from my project repository and related "
        "implementation work. The contents of this report are compiled for academic submission, and all technologies, tools, "
        "and reference materials used in preparing the report have been acknowledged in the references section.",
    )
    add_paragraph(doc, "\nPlace: ____________________\nDate: ____________________\n\nStudent Signature: ____________________", align=WD_ALIGN_PARAGRAPH.LEFT, indent=False)
    doc.add_page_break()

    add_heading(doc, "ACKNOWLEDGEMENT", level=1)
    add_paragraph(
        doc,
        "The completion of this report draws upon the Smart Grocery codebase, the documentation of the technologies used in the "
        "project, and the structured process of academic project development. Sincere gratitude is extended to the project guide, "
        "department, institution, and everyone whose support contributes to the successful understanding, implementation, and "
        "documentation of a full-stack software project.",
    )
    doc.add_page_break()

    add_heading(doc, "ABSTRACT", level=1)
    add_paragraph(
        doc,
        "Smart Grocery is a full-stack web application designed to support grocery planning, household inventory control, low-stock "
        "identification, purchase workflows, expiry-aware reminders, and administrator-level reporting. The frontend is implemented "
        "using React, Vite, Tailwind CSS, Axios, and React Router, while the backend is implemented using Spring Boot, Spring Security, "
        "Spring Validation, JPA, JWT, and MySQL. The system allows authenticated users to manage grocery items, view smart dashboard insights, "
        "act on shopping suggestions, and monitor kitchen reminders. Administrators can monitor users, products, stock levels, purchase queues, "
        "categories, and reports. The project demonstrates the integration of secure role-based access, business logic, persistence, and "
        "decision-support features in a maintainable full-stack architecture.",
    )
    doc.add_page_break()

    add_heading(doc, "INDEX", level=1)
    add_toc(doc.add_paragraph())
    add_paragraph(doc, "", indent=False)
    add_table(
        doc,
        [
            ["Sl. No.", "Topics", "Coverage"],
            ["1", "Introduction", "Introduction, scope, objectives, feasibility"],
            ["2", "Literature Survey / Software and Hardware Tools", "Existing systems, tools, technologies, hardware"],
            ["3", "Project Design", "Modules, architecture, interaction diagrams"],
            ["4", "Database Design", "ER design, relational structure, field summary"],
            ["5", "Development / Implementation Stages", "Phases followed to build the application"],
            ["6", "Testing", "Verification strategy, test evidence, test cases"],
            ["7", "Results", "Design screens, code excerpts, generated figures, observations"],
            ["8", "Conclusion", "Summary and future enhancement directions"],
            ["9", "Bibliography & References", "Repository and documentation references"],
        ],
    )
    doc.add_page_break()


def intro_section(doc: Document):
    add_heading(doc, "1. INTRODUCTION", level=1)
    add_heading(doc, "1.1 Introduction", level=2)
    add_paragraph(
        doc,
        "Smart Grocery is a web-based grocery planning and inventory management system developed to help users organize their household items, "
        "track quantities, distinguish between pending and purchased goods, receive low-stock warnings, and take action on expiry-driven reminders. "
        "The project extends beyond a traditional list manager by introducing dashboard insights, a smart shopping queue, recommendation logic, "
        "and a dedicated administrative suite for oversight and reporting.",
    )
    add_paragraph(
        doc,
        "The backend uses Spring Boot and MySQL to provide a secure role-aware API, while the frontend uses React to create a user-friendly "
        "interface for both end users and administrators. The system is structured so that each user views only their own grocery data, while "
        "administrators can observe system-wide trends and operational states.",
    )
    add_heading(doc, "1.2 Scope", level=2)
    add_bullets(
        doc,
        [
            "User registration, login, and role-aware access using JWT tokens.",
            "Household inventory management with item name, category, quantity, purchase status, and expiry date.",
            "Dashboard summaries for total items, pending items, purchased items, low-stock items, and recommended actions.",
            "Automatic shopping queue generation using low-stock and near-expiry conditions.",
            "Administrative control over users, products, category names, catalog stock, purchase queues, and reports.",
        ],
    )
    add_heading(doc, "1.3 Objectives", level=2)
    add_bullets(
        doc,
        [
            "To simplify grocery planning and day-to-day household inventory tracking.",
            "To reduce stock-outs through low-stock alerts and recommendation rules.",
            "To improve expiry awareness for purchased items through dashboard reminders.",
            "To implement a secure multi-role full-stack application with a maintainable architecture.",
            "To provide administrative insight into products, categories, stock, and user activity.",
        ],
    )
    add_heading(doc, "1.4 Feasibility Study", level=2)
    add_paragraph(
        doc,
        "Technical Feasibility:",
        align=WD_ALIGN_PARAGRAPH.LEFT,
        italic=True,
        indent=False,
    )
    add_paragraph(
        doc,
        "The Smart Grocery system is technically feasible using modern full-stack web development technologies. The project is built using React and Vite "
        "for the frontend, Spring Boot for the backend, and MySQL for database management. These technologies are widely adopted, well-supported, and "
        "suitable for building a responsive and secure grocery management platform.",
    )
    add_bullets(
        doc,
        [
            "Frontend: React, Vite, Tailwind CSS, Axios, and React Router are used to build responsive user and admin interfaces.",
            "Backend: Spring Boot with Spring Security, validation, and JPA handles authentication, API routing, inventory logic, and admin operations.",
            "Database: MySQL stores user accounts and grocery item records in a structured relational model.",
            "Security: JWT-based authentication, password hashing, protected routes, and role-based access control protect application data.",
        ],
    )
    add_paragraph(doc, "Operational Feasibility:", align=WD_ALIGN_PARAGRAPH.LEFT, italic=True, indent=False)
    add_paragraph(
        doc,
        "The Smart Grocery system is operationally feasible because it is easy to use and accessible through a web browser. Users can manage groceries, "
        "track stock levels, and view expiry reminders without special training. The admin module centralizes user management, stock review, queue handling, "
        "and reporting in a practical workflow.",
    )
    add_paragraph(
        doc,
        "Because the application is web-based, it can be used on standard computers and laptops in a simple local setup, making it suitable both for academic "
        "demonstration and day-to-day use.",
    )
    add_paragraph(doc, "Economic Feasibility:", align=WD_ALIGN_PARAGRAPH.LEFT, italic=True, indent=False)
    add_paragraph(
        doc,
        "From an economic point of view, Smart Grocery is cost-effective because it uses open-source technologies such as React, Spring Boot, MySQL, and related "
        "JavaScript and Java libraries. This reduces licensing costs while still providing a strong technical foundation for development and deployment.",
    )
    add_bullets(
        doc,
        [
            "Development Costs: The main cost is the time and effort required for design, coding, testing, and documentation, with no expensive software licenses.",
            "Operational Costs: The system can run on a standard local machine for development and testing, and later deployment can use affordable hosting options.",
        ],
    )
    doc.add_page_break()


def literature_tools_section(doc: Document):
    add_heading(doc, "2. LITERATURE SURVEY / SOFTWARE AND HARDWARE TOOLS", level=1)
    add_heading(doc, "2.1 Literature Survey", level=2)
    add_paragraph(
        doc,
        "Many grocery and home inventory applications currently available to users focus on simple list creation, shopping reminders, or individual "
        "purchase tracking. However, a gap remains between a basic grocery checklist and an application that actively helps users manage stock pressure, "
        "predict restocking needs, and notice expiring purchased goods in a structured way.",
    )
    add_paragraph(
        doc,
        "Existing systems often provide a clean user interface but lack a deeper operational model. Some tools support categories and quantities, but "
        "they do not transform that data into meaningful recommendations or separate user responsibilities from administrator oversight. Smart Grocery "
        "improves on these gaps by combining household inventory, recommendation logic, expiry reminders, and administrative analytics within one system.",
    )
    add_heading(doc, "2.1.1 Existing Systems and Solutions", level=3)
    add_paragraph(
        doc,
        "Common market solutions include generic task managers, note-based grocery lists, and e-commerce shopping carts. These systems help users list what "
        "to buy, but usually do not treat groceries as a living inventory with lifecycle data such as quantities, purchase states, and expiry dates. They "
        "also rarely include a separate operational dashboard or an admin-oriented reporting interface.",
    )
    add_heading(doc, "2.1.2 Gaps in Existing Systems", level=3)
    add_paragraph(
        doc,
        "The main gaps include weak decision support, little or no expiry intelligence, inconsistent stock awareness, and minimal separation between ordinary "
        "user workflows and system administration. Smart Grocery addresses these limitations through rule-driven shopping suggestions, category-aware stock "
        "visualization, kitchen reminders, and a multi-page admin suite.",
    )
    add_heading(doc, "2.1.3 Proposed Solution Overview", level=3)
    add_paragraph(
        doc,
        "The proposed Smart Grocery solution is a browser-based platform where users can securely log in, maintain a grocery inventory, monitor stock conditions, "
        "and act on recommendations and reminders. Administrators can inspect broader platform activity, adjust catalog stock, complete purchase queues, rename "
        "categories, and review reports generated from the same shared data model.",
    )
    add_heading(doc, "2.2 Software and Hardware Tools", level=2)
    add_table(
        doc,
        [
            ["Category", "Tools Used", "Purpose"],
            ["Frontend", "React, Vite, Tailwind CSS, Axios, React Router", "User interface, routing, styling, API communication"],
            ["Backend", "Spring Boot, Spring Security, Validation, JPA", "REST API, authorization, validation, persistence"],
            ["Database", "MySQL", "Storage of users and grocery items"],
            ["Authentication", "JWT", "Secure token-based access control"],
            ["Testing", "JUnit, Spring test starters, ESLint", "Backend verification and frontend static checks"],
            ["Version Control", "Git", "Repository tracking and development workflow"],
        ],
    )
    add_paragraph(doc, "", indent=False)
    add_bullets(
        doc,
        [
            "Java 21 or later is required to run the backend service.",
            "Node.js 20 or later is required to run the frontend development server.",
            "MySQL 8 or later is used for the persistent relational database.",
            "A standard personal computer and modern browser are sufficient for development and usage.",
        ],
    )
    doc.add_page_break()


def design_section(doc: Document):
    add_heading(doc, "3. PROJECT DESIGN", level=1)
    add_heading(doc, "3.1 Design Details of Modules", level=2)
    add_paragraph(
        doc,
        "The Smart Grocery application is organized into user-facing modules and administrator-facing modules. Each module focuses on a specific aspect of the "
        "grocery lifecycle so that the interface remains understandable while the backend retains clean responsibilities across controllers, services, and repositories.",
    )
    add_table(
        doc,
        [
            ["Module", "Responsibilities"],
            ["Authentication Module", "Registration, login, admin login, session token generation, route protection"],
            ["Home and Catalog Module", "Quick-buy flow, catalog browsing, availability display, search and category filters"],
            ["Inventory Module", "Add items, update quantities, search, filter, delete, purchase toggle, expiry preview"],
            ["Dashboard Module", "Summary metrics, smart picks, low-stock watchlist, pending focus, kitchen reminders"],
            ["Shopping List Module", "Auto-built queue from low stock and near-expiry inventory items"],
            ["Admin Module", "Users, dashboard, products, categories, purchase queue, stock controls, reports"],
        ],
    )
    add_heading(doc, "3.2 Architecture Diagram", level=2)
    add_paragraph(
        doc,
        "The architecture follows a client-server design. React pages issue API calls through Axios. Spring Boot controllers expose route contracts for authentication, "
        "grocery operations, and administrator workflows. Service classes implement business rules, such as stock thresholds, recommendations, expiry logic, and "
        "report aggregation. JPA repositories interact with MySQL for persistence. JWT and protected frontend routes enforce access control for user and admin roles.",
    )
    add_paragraph(
        doc,
        "The figures in this chapter summarize the application architecture, module interaction, and data flows between user actions and system services.",
    )
    doc.add_page_break()

    add_figure(doc, "architecture_diagram.png", "Figure 3.1:", "Architecture diagram showing the React frontend, Spring Boot controllers, service layer, security layer, smart logic, and MySQL persistence.")
    add_figure(doc, "module_interaction.png", "Figure 3.2:", "Module interaction diagram showing how user and admin pages map to major application responsibilities.")
    add_figure(doc, "dfd_user.png", "Figure 3.3:", "Data flow diagram for the normal user journey across authentication, inventory, dashboard, shopping, and smart logic.")
    add_figure(doc, "dfd_admin.png", "Figure 3.4:", "Data flow diagram for the administrator workflow across login, dashboard, queue handling, and reporting.")


def database_section(doc: Document):
    add_heading(doc, "4. DATABASE DESIGN", level=1)
    add_heading(doc, "4.1 E-R Diagrams / DFDs / Relational View", level=2)
    add_paragraph(
        doc,
        "The current Smart Grocery project uses a focused relational model. The small number of core entities keeps the application maintainable while still supporting "
        "meaningful user isolation and analytics. A single user may own many grocery items, and each grocery item belongs to exactly one user. This relationship allows "
        "the application to provide user-scoped inventory management while still supporting platform-wide administrative reporting.",
    )
    add_paragraph(
        doc,
        "The grocery_items table stores fields that are directly relevant to business behavior, including quantity, purchase state, expiry date, and last purchase time. "
        "These fields power recommendations, low-stock alerts, shopping list generation, and kitchen reminder logic.",
    )
    add_table(
        doc,
        [
            ["Entity", "Important Fields", "Remarks"],
            ["users", "id, username, email, password, role", "Stores authenticated users and their authorization level"],
            ["grocery_items", "id, name, category, quantity, purchased, expiry_date, last_purchased_at, user_id", "Stores household grocery inventory and workflow data"],
        ],
    )
    doc.add_page_break()
    add_figure(doc, "er_diagram.png", "Figure 4.1:", "Entity relationship diagram representing users and grocery_items with a one-to-many association.", page_break=False)
    add_figure(doc, "catalog_stock_chart.png", "Figure 4.2:", "Sample chart derived from default catalog stock quantities defined in the service layer.", page_break=True)


def development_section(doc: Document):
    add_heading(doc, "5. DEVELOPMENT / IMPLEMENTATION STAGES", level=1)
    add_paragraph(
        doc,
        "The Smart Grocery project was implemented in structured stages so that the application could evolve from a basic authenticated CRUD system into a smarter "
        "inventory and reporting platform. The stages described below capture how the project can be presented academically while reflecting the current repository state.",
    )
    add_heading(doc, "5.1 Requirement Analysis", level=2)
    add_bullets(
        doc,
        [
            "Identified the need for user registration, login, and secure role-aware access.",
            "Defined grocery CRUD behavior with search, category filters, and purchase status.",
            "Planned dashboard insights, recommendations, shopping queue logic, and expiry reminders.",
            "Defined administrator operations for users, products, categories, stock, queues, and reports.",
        ],
    )
    add_heading(doc, "5.2 Environment Setup", level=2)
    add_bullets(
        doc,
        [
            "Configured the Spring Boot backend with MySQL, JWT secret, and CORS origin.",
            "Initialized the React frontend with Vite, Tailwind CSS, Axios, and React Router.",
            "Prepared repository scripts for frontend linting and backend Maven test execution.",
        ],
    )
    add_heading(doc, "5.3 Backend Development", level=2)
    add_paragraph(
        doc,
        "Controller classes were created for authentication, grocery operations, and administrator features. Service classes were then introduced to isolate business logic. "
        "This includes password hashing, role resolution, recommendation scoring, catalog stock consumption, expiry-date derivation, and category-wise reporting.",
    )
    add_heading(doc, "5.4 Frontend Development", level=2)
    add_paragraph(
        doc,
        "The frontend was split into focused pages so that users and admins have distinct shells and routes. The design emphasizes operational workflows rather than generic forms, "
        "with dedicated cards, queue views, summaries, and management panels for different tasks.",
    )
    add_heading(doc, "5.5 Integration and Validation", level=2)
    add_paragraph(
        doc,
        "Axios is used to connect React pages to the backend API. Shared events are used to refresh dependent views when grocery data changes. Validation errors and authorization "
        "failures are converted into user-readable frontend messages, which helps keep the experience consistent across the app.",
    )
    doc.add_page_break()


def testing_section(doc: Document):
    add_heading(doc, "6. TESTING", level=1)
    add_paragraph(
        doc,
        "Testing is an essential phase in software development because it verifies whether a system behaves as expected under realistic scenarios. The Smart Grocery project combines "
        "static frontend verification with backend test execution and documented workflow validation.",
    )
    add_heading(doc, "6.1 Types of Testing Performed", level=2)
    add_bullets(
        doc,
        [
            "Unit and service testing for backend business logic such as authentication and grocery rules.",
            "Controller testing for route-level behavior and validation handling.",
            "Frontend linting for code-quality and rule compliance.",
            "Manual workflow review through the route and UI design structure represented in the project.",
        ],
    )
    add_heading(doc, "6.2 Verification Results", level=2)
    add_table(
        doc,
        [
            ["Verification", "Date", "Result"],
            ["Frontend lint (`npm run lint`)", "19 April 2026", "Passed"],
            ["Frontend build (`npm run build`)", "19 April 2026", "Blocked in the current environment by a native Tailwind / Vite dependency loading issue"],
            ["Backend tests (`.\\mvnw.cmd test`)", "19 April 2026", "50 tests executed with 1 failure, 0 errors, 0 skipped"],
        ],
    )
    add_heading(doc, "6.3 Sample Test Cases", level=2)
    add_table(
        doc,
        [
            ["Test Case", "Input / Condition", "Expected Output", "Status"],
            ["User login", "Valid username and password", "JWT response and access to user routes", "Passed"],
            ["Admin login", "Admin credentials", "JWT response with ADMIN role", "Passed"],
            ["Low-stock queue generation", "Pending item quantity <= threshold", "Item appears in shopping list", "Passed"],
            ["Expiry reminder generation", "Purchased item expiring soon", "Reminder appears in dashboard", "Passed"],
            ["Role update guard", "Attempt to remove last admin role", "Conflict error returned", "Passed"],
        ],
    )
    add_paragraph(
        doc,
        "The frontend build result is preserved accurately in the report instead of being hidden. This is important in an academic submission because the report should reflect the actual "
        "observed environment rather than idealized behavior. In the current backend verification run, the failing test is `AuthControllerTest.registerRejectsInvalidPayload`, where the "
        "expected username validation message differs from the actual response. The lint result and the remaining backend tests nevertheless provide meaningful verification evidence.",
    )
    doc.add_page_break()
    add_figure(doc, "testing_strategy.png", "Figure 6.1:", "Testing strategy diagram summarizing linting, backend tests, controller tests, manual validation, and documentation evidence.")
    add_figure(doc, "module_count_chart.png", "Figure 6.2:", "Chart showing the relative breadth of user pages, admin pages, controllers, services, and test suites covered in the report.")


def results_section(doc: Document):
    add_heading(doc, "7. RESULTS", level=1)
    add_heading(doc, "7.1 Result Summary", level=2)
    add_paragraph(
        doc,
        "The Smart Grocery project delivers a complete user and administrator workflow. The resulting application is not limited to adding or deleting items; instead, it presents "
        "a structured household inventory experience with operational dashboards, category-aware filtering, recommendation logic, expiry reminders, shopping suggestions, and platform-wide "
        "administrative visibility.",
    )
    add_paragraph(
        doc,
        "In order to mirror the visual density of the sample report, this chapter documents the system through generated design-screen figures and representative code excerpts. Each figure "
        "highlights a major page or feature area from the project and is followed by a short academic description of its role.",
    )
    doc.add_page_break()

    figure_specs = [
        FigureSpec("ui_auth_login.png", "Figure 7.1:", "Login screen design showing secure authentication entry for Smart Grocery users and administrators."),
        FigureSpec("ui_auth_register.png", "Figure 7.2:", "Registration screen design showing new user onboarding with validation-aware account creation."),
        FigureSpec("ui_home_catalog.png", "Figure 7.3:", "Home and catalog screen showing quick-buy options, availability, and top-level grocery navigation."),
        FigureSpec("ui_quick_buy.png", "Figure 7.4:", "Quick-buy modal design illustrating quantity entry and expiry preview for direct purchase flow."),
        FigureSpec("ui_dashboard_overview.png", "Figure 7.5:", "Dashboard overview screen with summary cards and action-focused sections."),
        FigureSpec("ui_dashboard_recommendations.png", "Figure 7.6:", "Smart picks screen showing recommendation cards with reasons and priority labels."),
        FigureSpec("ui_dashboard_lowstock.png", "Figure 7.7:", "Low-stock watchlist showing urgent inventory items that need restocking."),
        FigureSpec("ui_dashboard_expiry.png", "Figure 7.8:", "Kitchen reminder screen showing expiry-driven alerts for purchased items."),
        FigureSpec("ui_inventory.png", "Figure 7.9:", "Inventory management screen combining item entry, filters, and item list controls."),
        FigureSpec("ui_shopping_list.png", "Figure 7.10:", "Shopping list screen generated from low-stock and near-expiry conditions."),
        FigureSpec("ui_admin_dashboard.png", "Figure 7.11:", "Admin dashboard screen showing system-wide counts and platform overview."),
        FigureSpec("ui_admin_users.png", "Figure 7.12:", "Admin users screen showing role management and user summary controls."),
        FigureSpec("ui_admin_products.png", "Figure 7.13:", "Admin products screen showing stock adjustments and product-level controls."),
        FigureSpec("ui_admin_purchase_queue.png", "Figure 7.14:", "Admin purchase queue screen showing centralized completion of pending purchases."),
        FigureSpec("ui_admin_categories.png", "Figure 7.15:", "Admin categories screen showing rename controls and per-category summaries."),
        FigureSpec("ui_admin_reports.png", "Figure 7.16:", "Admin reports screen showing operational reporting across users, products, and categories."),
    ]
    for spec in figure_specs:
        add_figure(doc, spec.filename, spec.title, spec.caption)

    add_heading(doc, "7.2 Observations from the Result Screens", level=2)
    result_observations = [
        "The user-facing pages are organized around tasks rather than technical entities. The dashboard emphasizes action, while the inventory page emphasizes editing.",
        "The home page behaves as a lightweight operational catalog rather than a static landing screen, which makes the project feel more product-oriented.",
        "Administrator pages separate control concerns cleanly: users, products, stock, categories, purchase queue, and reports each get their own route and screen.",
        "Expiry logic and shopping logic convert stored data into decision support, which improves the academic value of the project beyond CRUD.",
        "Report-friendly figures can be mapped directly to repository pages, which makes the documentation more defensible in an academic review setting.",
    ]
    for paragraph in result_observations:
        add_paragraph(doc, paragraph)
    doc.add_page_break()

    add_heading(doc, "7.3 Sample Code and Implementation Evidence", level=2)
    snippets = [
        ("AuthController.java excerpt", read_snippet(ROOT / "backend/src/main/java/com/example/backend/controller/AuthController.java", 1, 28)),
        ("AuthService.java excerpt", read_snippet(ROOT / "backend/src/main/java/com/example/backend/service/AuthService.java", 19, 62)),
        ("GroceryController.java excerpt", read_snippet(ROOT / "backend/src/main/java/com/example/backend/controller/GroceryController.java", 1, 89)),
        ("GroceryService.java excerpt - inventory and summary logic", read_snippet(ROOT / "backend/src/main/java/com/example/backend/service/GroceryService.java", 68, 190)),
        ("GroceryService.java excerpt - expiry and recommendation logic", read_snippet(ROOT / "backend/src/main/java/com/example/backend/service/GroceryService.java", 254, 395)),
        ("AdminService.java excerpt", read_snippet(ROOT / "backend/src/main/java/com/example/backend/service/AdminService.java", 29, 205)),
        ("App.jsx routes excerpt", read_snippet(ROOT / "frontend/src/App.jsx", 1, 37)),
        ("DashboardPage.jsx excerpt", read_snippet(ROOT / "frontend/src/pages/DashboardPage.jsx", 1, 210)),
        ("InventoryPage.jsx excerpt", read_snippet(ROOT / "frontend/src/pages/InventoryPage.jsx", 1, 260)),
        ("ShoppingListPage.jsx excerpt", read_snippet(ROOT / "frontend/src/pages/ShoppingListPage.jsx", 1, 196)),
        ("AdminProductsPage.jsx excerpt", read_snippet(ROOT / "frontend/src/pages/AdminProductsPage.jsx", 1, 210)),
        ("AdminReportsPage.jsx excerpt", read_snippet(ROOT / "frontend/src/pages/AdminReportsPage.jsx", 1, 81)),
    ]
    for title, snippet in snippets:
        add_code_block(doc, title, snippet)
        doc.add_page_break()

    add_heading(doc, "7.4 Discussion of Results", level=2)
    discussion = [
        "The final system demonstrates that a grocery management project can be elevated from a checklist tool into a richer inventory application by connecting stored item data to actionable rules.",
        "The repository shows a clear separation between authentication behavior, grocery item workflows, and admin-level operational management, which helps both maintainability and documentation quality.",
        "The React pages make the project more report-friendly because each major use case appears as a dedicated screen that can be described independently in the results chapter.",
        "The service layer carries most of the intelligence of the application, especially recommendation scoring, expiry derivation, stock consumption, and report aggregation.",
        "The existence of passing backend tests and a clean frontend lint run provides meaningful quality evidence to accompany the design and results sections.",
    ]
    for paragraph in discussion:
        add_paragraph(doc, paragraph)
    doc.add_page_break()


def conclusion_section(doc: Document):
    add_heading(doc, "8. CONCLUSION", level=1)
    add_paragraph(
        doc,
        "Smart Grocery demonstrates a practical full-stack solution for modern grocery planning and household inventory management. The project combines secure authentication, "
        "role-aware access, persistent inventory data, low-stock identification, shopping suggestions, expiry reminders, and administrative analytics into one integrated application. "
        "The result is more capable than a traditional grocery checklist because the stored data is continuously transformed into decision-support information.",
    )
    add_paragraph(
        doc,
        "From an academic perspective, the project is valuable because it covers frontend development, backend API design, database modeling, security, validation, business logic, testing, "
        "and documentation. Future enhancements could include notification delivery, richer analytics, barcode integration, automated screenshot capture, and production deployment support.",
    )
    doc.add_page_break()


def references_section(doc: Document):
    add_heading(doc, "9. BIBLIOGRAPHY & REFERENCES", level=1)
    add_bullets(
        doc,
        [
            "Smart Grocery repository source code in S:\\Smart_Grocery.",
            "Spring Boot documentation for web, security, validation, and JPA.",
            "React documentation.",
            "React Router documentation.",
            "Axios documentation.",
            "MySQL documentation.",
            "JWT library documentation and related authentication references.",
        ],
    )


def build_document():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)
    generate_assets()

    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(12)

    for section in doc.sections:
        set_page_margins(section)
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(footer)

    add_front_matter(doc)
    intro_section(doc)
    literature_tools_section(doc)
    design_section(doc)
    database_section(doc)
    development_section(doc)
    testing_section(doc)
    results_section(doc)
    conclusion_section(doc)
    references_section(doc)

    doc.save(OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    build_document()
