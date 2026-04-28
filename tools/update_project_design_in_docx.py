# pyright: reportMissingImports=false, reportMissingTypeStubs=false, reportUnknownMemberType=false, reportPrivateUsage=false
from __future__ import annotations

from pathlib import Path
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE_PATH = ROOT / "docs" / "Smart_Grocery_Project_Report.docx"
FALLBACK_SOURCE = ROOT / "docs" / "Smart_Grocery_Project_Report_Feasibility_Updated.docx"
OUTPUT_PATH = ROOT / "docs" / "Smart_Grocery_Project_Report.docx"
TEMP_PATH = ROOT / "docs" / "Smart_Grocery_Project_Report_Temp_Project_Design.docx"


def remove_paragraph(paragraph):
    element = paragraph._element
    parent = element.getparent()
    if parent is None:
        return
    parent.remove(element)


def insert_paragraph_after(paragraph, text: str = ""):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    new_paragraph = paragraph._parent.add_paragraph()
    new_paragraph._p = new_p
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def set_run_font(run, size=12, bold=False, italic=False, color: RGBColor | None = None, name="Times New Roman"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def format_body(paragraph, text: str):
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.first_line_indent = Inches(0.4)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def format_module_title(paragraph, text: str):
    paragraph.text = ""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.line_spacing = 1.2
    paragraph.paragraph_format.space_before = Pt(8)
    paragraph.paragraph_format.space_after = Pt(2)
    paragraph.paragraph_format.first_line_indent = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=12, italic=True, color=RGBColor(79, 129, 189))


def format_bullet(paragraph, text: str):
    paragraph.text = ""
    paragraph.style = "List Bullet"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def format_sub_bullet(paragraph, text: str):
    paragraph.text = ""
    paragraph.style = "List Bullet 2"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    paragraph.paragraph_format.line_spacing = 1.5
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.space_before = Pt(0)
    run = paragraph.add_run(text)
    set_run_font(run, size=12)


def resolve_source() -> Path:
    if SOURCE_PATH.exists():
        return SOURCE_PATH
    if FALLBACK_SOURCE.exists():
        return FALLBACK_SOURCE
    raise FileNotFoundError("No Smart Grocery report document found.")


def main():
    source = resolve_source()
    shutil.copy2(source, TEMP_PATH)
    doc = Document(TEMP_PATH)
    paragraphs = doc.paragraphs

    design_heading_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "3.1 Design Details of Modules")
    arch_heading_idx = next(i for i, p in enumerate(paragraphs) if p.text.strip() == "3.2 Architecture Diagram")

    for idx in range(arch_heading_idx - 1, design_heading_idx, -1):
        remove_paragraph(doc.paragraphs[idx])

    anchor = doc.paragraphs[design_heading_idx]
    blocks = [
        ("body", "In this section, we will outline the project design, including the design details of the system modules and the architecture of the application. This section aims to provide a clear understanding of how the system is structured, both in terms of the backend and frontend, and how various components interact."),
        ("body", "The Smart Grocery system is divided into several core modules. Each module serves a specific function to ensure the smooth operation of the system. The modules are as follows:"),

        ("module", "1. User Authentication Module"),
        ("bullet", "Functionality: This module handles user login, registration, and role-aware access management."),
        ("bullet", "Components:"),
        ("subbullet", "Login Form: Allows users to enter their credentials using username or email and password."),
        ("subbullet", "Registration Form: New users can register for the platform by providing username, email, and password."),
        ("subbullet", "Token Management: Once a user logs in, a JWT token is issued and used for protected API access."),
        ("subbullet", "Role Validation: The system distinguishes between normal users and administrators for route access."),
        ("bullet", "Technology: Spring Security, JWT, password hashing, React auth pages, and protected routes."),

        ("module", "2. Home and Catalog Module"),
        ("bullet", "Functionality: This module allows users to browse grocery catalog items, check availability, and quickly purchase items into their inventory."),
        ("bullet", "Components:"),
        ("subbullet", "Catalog Cards: Display item name, category, price, stock quantity, and availability state."),
        ("subbullet", "Search and Filter: Allows users to filter catalog items by category and search keyword."),
        ("subbullet", "Quick Buy Modal: Lets users purchase a catalog item directly with quantity and expiry preview."),
        ("subbullet", "Home Summary: Shows pending, purchased, and currently available grocery counts."),
        ("bullet", "Technology: React components, Axios API calls, dynamic filters, and backend catalog endpoints."),

        ("module", "3. Inventory Management Module"),
        ("bullet", "Functionality: Manages grocery item creation, listing, updates, deletion, and purchased-state tracking."),
        ("bullet", "Components:"),
        ("subbullet", "Add Item Form: Allows users to enter item name, category, quantity, purchased state, and expected expiry."),
        ("subbullet", "Inventory List View: Displays grocery items with quantity, category, status, and timestamps."),
        ("subbullet", "Filter Controls: Supports search, category filtering, and purchased / pending filtering."),
        ("subbullet", "Item Actions: Users can mark items as bought, move them back to pending, or delete them."),
        ("bullet", "Technology: GroceryController routes, GroceryService logic, JPA persistence, and React list management."),

        ("module", "4. Dashboard and Recommendation Module"),
        ("bullet", "Functionality: Provides users with an overview of grocery metrics, low-stock pressure, recommendations, and recent activity."),
        ("bullet", "Components:"),
        ("subbullet", "Summary Cards: Show total items, pending items, purchased items, and low-stock items."),
        ("subbullet", "Smart Picks: Displays recommendation cards derived from prior purchase and low-stock conditions."),
        ("subbullet", "Low-Stock Watchlist: Highlights pending items that require restocking soon."),
        ("subbullet", "Recent Activity: Lists recent grocery actions and the current pending focus."),
        ("bullet", "Technology: Aggregated API calls, recommendation scoring, and dashboard-specific React sections."),

        ("module", "5. Shopping List and Expiry Reminder Module"),
        ("bullet", "Functionality: Handles automatic shopping suggestions and expiry-aware kitchen reminders for purchased items."),
        ("bullet", "Components:"),
        ("subbullet", "Shopping Queue: Builds a buy list from low-stock and near-expiry conditions."),
        ("subbullet", "Priority Labels: Assigns HIGH, MEDIUM, or LOW priority to shopping suggestions."),
        ("subbullet", "Expiry Alerts: Shows items expiring soon, expiring today, or already overdue."),
        ("subbullet", "Acknowledge Action: Allows users to remove an expiry reminder by deleting the expired item."),
        ("bullet", "Technology: Expiry-date rules, shopping-list logic, dashboard reminders, and user-scoped API endpoints."),

        ("module", "6. Admin Management and Reporting Module"),
        ("bullet", "Functionality: The admin module allows the administrator to manage users, products, categories, purchase queues, stock levels, and reports."),
        ("bullet", "Components:"),
        ("subbullet", "Admin Dashboard: Displays total users, products, pending items, purchased items, categories, and low-stock counts."),
        ("subbullet", "User Management: Admins can view users, update roles, and delete accounts with admin-safety rules."),
        ("subbullet", "Product and Stock Management: Admins can edit product quantities, delete products, and restock catalog items."),
        ("subbullet", "Purchase Queue and Reports: Admins can complete pending purchases and view category-wise reports and breakdowns."),
        ("bullet", "Technology: AdminController routes, AdminService aggregation logic, catalog stock state, and reporting endpoints."),
    ]

    current = anchor
    for kind, text in blocks:
        current = insert_paragraph_after(current)
        if kind == "body":
            format_body(current, text)
        elif kind == "module":
            format_module_title(current, text)
        elif kind == "subbullet":
            format_sub_bullet(current, text)
        else:
            format_bullet(current, text)

    doc.save(TEMP_PATH)
    shutil.copy2(TEMP_PATH, OUTPUT_PATH)
    print(OUTPUT_PATH)


if __name__ == "__main__":
    main()
